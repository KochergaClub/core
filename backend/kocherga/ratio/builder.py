import logging

logger = logging.getLogger(__name__)

from pathlib import Path
import re
import subprocess
import os
from tempfile import TemporaryDirectory
import contextlib

from kocherga.chrome import get_browser

import docx

import kocherga.gdrive

from .content import Section

SECTIONS = [
    {"id": "1.1", "name": "Введение, день 1", "handbook": False},
    {"id": "1.2", "name": "Факторизация целей"},
    {"id": "1.3", "name": "Внутренний симулятор"},
    {"id": "1.4", "name": "Премортемы"},
    {"id": "1.5", "name": "Ожидаемая полезность"},
    {"id": "2.1", "name": "Введение, день 2", "handbook": False},
    {"id": "2.2", "name": "Установка навыков"},
    {"id": "2.3", "name": "Факторизация избеганий"},
    {"id": "2.4", "name": "Внутренние конфликты"},
    {"id": "2.5", "name": "Оценки вероятностей"},
]


def events_root_id():
    # TODO - move to config
    return "1WeeuH0Rtp8ODmJ2vbUKE5FklFkZ96Qn4"


def workshops_root_id():
    return kocherga.gdrive.parent_id(events_root_id())


def humanized_event_date(event_name):
    match = re.match(r'(\d+)-(\d+) воркшоп$', event_name)
    (year, month) = match.groups()
    month_names = 'январь февраль март апрель май июнь июль август сентябрь октябрь ноябрь декабрь'.split(
        ' '
    )
    month_name = month_names[int(month) - 1]
    return f'{month_name} {year}'


# TODO - fetch from google drive via API or make more customizable
@contextlib.contextmanager
def serve_slides(
    path=os.environ["HOME"] + "/Google Drive/Rationality/Workshops/Секции",
):
    server = subprocess.Popen(["python3", "-m", "http.server"], cwd=path)
    try:
        yield server
    finally:
        server.kill()


def prepare_folders(event_name):
    logger.info("Preparing folders")
    event_folder_id = kocherga.gdrive.create_folder(events_root_id(), event_name)
    content_folder_id = kocherga.gdrive.create_folder(event_folder_id, "Пострассылка")
    public_folder_id = kocherga.gdrive.create_folder(
        content_folder_id, "Материалы по прикладной рациональности"
    )
    return public_folder_id


async def build_slides(folder_id, sections=SECTIONS):
    logger.info("Building slides")

    with serve_slides(), TemporaryDirectory() as tmp_dir:
        slides_folder_id = kocherga.gdrive.create_folder(folder_id, "Слайды")

        async with get_browser() as browser:
            page = await browser.newPage()

            for section in sections:
                name = f'{section["id"]} - {section["name"]}.pdf'

                if kocherga.gdrive.find_in_folder(
                    slides_folder_id, name, missing_ok=True
                ):
                    # file already exists
                    continue

                logger.info(f'Section {section["name"]}')
                await page.goto("http://localhost:8000/" + section["name"])

                pdf_bytes = await page.pdf(width="32.04cm", height="18.03cm")

                tmp_filename = str(Path(tmp_dir) / section["id"])
                with open(tmp_filename, 'wb') as fh:
                    fh.write(pdf_bytes)

                kocherga.gdrive.upload_file(tmp_filename, name, slides_folder_id)


def build_handbooks(event_name, sections=SECTIONS):
    # TODO - get from gdrive
    DOCX_TEMPLATE = '/Users/berekuk/Google Drive/Rationality/Workshops/Шаблоны/Стили рабочих тетрадей.docx'

    for day in '1', '2':
        document = docx.Document(DOCX_TEMPLATE)

        document.add_heading(
            'Прикладная рациональность: рабочая тетрадь'
        ).style = 'Заголовок тетради'
        # TODO - add "День 1" / "День 2" subheader

        document.add_heading(
            f'Воркшоп по прикладной рациональности, Москва\n{humanized_event_date(event_name)}'
        ).style = 'Метаданные мероприятия'

        # TODO - add TOC

        for section_info in sections:
            if not section_info['id'].startswith(day + '.'):
                continue
            if not section_info.get('handbook', True):
                continue
            section = Section(section_info['name'])
            logger.info(section.name)

            document.add_heading(section.name).style = 'Heading 1'
            section.populate_docx(document)

        # TODO - add glossary

        # TODO - save to gdrive
        document.save(f'/Users/berekuk/Downloads/Handbook {day}.docx')


def build_extra_files(folder_id):
    sources_folder_id = kocherga.gdrive.parent_id(folder_id)

    # # schedule
    # # TODO - build schedule from html to pdf
    # event_folder_id = kocherga.gdrive.parent_id(sources_folder_id)
    # schedule_gdoc_id = kocherga.gdrive.find_in_folder(event_folder_id, "Расписание")
    # filename = kocherga.gdrive.gdoc_to_pdf(schedule_gdoc_id)
    # kocherga.gdrive.upload_file(filename, "Расписание.pdf", folder_id)

    # additional materials
    materials_gdoc_id = kocherga.gdrive.find_in_folder(
        sources_folder_id, "Дополнительные материалы"
    )
    filename = kocherga.gdrive.gdoc_to_pdf(materials_gdoc_id)
    kocherga.gdrive.upload_file(filename, "Материалы.pdf", folder_id)

    # feedback form
    form_filename = 'Впечатления от воркшопа'  # TODO - name according to the event
    if kocherga.gdrive.find_in_folder(
        sources_folder_id, form_filename, missing_ok=True
    ):
        # already exists
        return
    templates_folder_id = kocherga.gdrive.find_in_folder(workshops_root_id(), 'Шаблоны')
    post_form_id = kocherga.gdrive.find_in_folder(
        templates_folder_id, 'Впечатления от воркшопа'
    )

    kocherga.gdrive.gdrive().files().copy(
        fileId=post_form_id,
        body={"name": form_filename, "parents": [sources_folder_id]},
    ).execute()


async def build(event_name):
    folder_id = prepare_folders(event_name)

    build_extra_files(folder_id)
    await build_slides(folder_id)
    build_handbooks(event_name)

    kocherga.gdrive.share_to_public(folder_id)
